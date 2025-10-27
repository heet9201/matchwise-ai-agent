from fastapi import UploadFile, HTTPException, status
import PyPDF2
from docx import Document
import io
import logging
import httpx
import re
from typing import Set, Optional, Dict
from datetime import datetime
from bs4 import BeautifulSoup
from urllib.parse import urlparse

logger = logging.getLogger(__name__)

class FileProcessingError(Exception):
    """Custom exception for file processing errors"""
    pass

class FileService:
    SUPPORTED_FORMATS: Set[str] = {'.pdf', '.doc', '.docx'}
    MAX_FILE_SIZE: int = 10 * 1024 * 1024  # 10MB

    async def process_file(self, file: UploadFile) -> str:
        """
        Process uploaded PDF or DOCX files and extract text content
        
        Args:
            file: The uploaded file to process
            
        Returns:
            str: Extracted text content from the file
            
        Raises:
            FileProcessingError: If there's an error processing the file
            HTTPException: If the file format is unsupported or file is too large
        """
        try:
            # Validate file format
            file_ext = self._get_file_extension(file.filename)
            if file_ext not in self.SUPPORTED_FORMATS:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Unsupported file format. Supported formats: {', '.join(self.SUPPORTED_FORMATS)}"
                )

            # Read and validate file content
            content = await file.read()
            if len(content) > self.MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File size exceeds maximum limit of {self.MAX_FILE_SIZE // 1024 // 1024}MB"
                )

            # Process file based on format with progress tracking
            try:
                start_time = datetime.now()
                logger.info(f"Started processing file: {file.filename}")

                if file_ext == '.pdf':
                    text = self._extract_pdf_text(content)
                else:  # .doc or .docx
                    text = self._extract_docx_text(content)

                # Validate extracted content
                if not text or len(text.strip()) < 50:
                    raise FileProcessingError("Extracted text is too short or empty")
                
                processing_time = (datetime.now() - start_time).total_seconds()
                logger.info(f"Completed processing {file.filename} in {processing_time:.2f} seconds")

                return text

            except FileProcessingError as e:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=str(e)
                )
            except Exception as e:
                logger.error(f"Error processing file {file.filename}: {str(e)}", exc_info=True)
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Error processing file"
                )
        except Exception as e:
            logger.error(f"Error reading file {file.filename}: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Error reading file"
            )

    def _get_file_extension(self, filename: str) -> str:
        """Extract and validate file extension"""
        if not filename or '.' not in filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid filename"
            )
        return filename[filename.rindex('.'):].lower()

    def _extract_pdf_text(self, content: bytes) -> str:
        """
        Extract text from PDF file
        
        Args:
            content: Raw bytes of the PDF file
            
        Returns:
            str: Extracted text content
            
        Raises:
            FileProcessingError: If text extraction fails
        """
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                raise FileProcessingError("PDF file is empty")

            text = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)

            return "\n".join(text)

        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}", exc_info=True)
            raise FileProcessingError(f"Failed to extract text from PDF: {str(e)}")

    def _extract_docx_text(self, content: bytes) -> str:
        """
        Extract text from DOCX file
        
        Args:
            content: Raw bytes of the DOCX file
            
        Returns:
            str: Extracted text content
            
        Raises:
            FileProcessingError: If text extraction fails
        """
        try:
            doc = Document(io.BytesIO(content))
            
            if len(doc.paragraphs) == 0:
                raise FileProcessingError("DOCX file is empty")

            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            return "\n".join(text)

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}", exc_info=True)
            raise FileProcessingError(f"Failed to extract text from DOCX: {str(e)}")

    async def fetch_job_description_from_url(self, url: str) -> Dict[str, str]:
        """
        Fetch and extract job description content from a URL
        
        Args:
            url: The URL of the job posting
            
        Returns:
            dict: Contains 'text' (extracted content) and 'company_name' (if found)
            
        Raises:
            HTTPException: If URL is invalid or content cannot be fetched
        """
        try:
            # Validate URL
            if not self._is_valid_url(url):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Invalid URL format"
                )
            
            logger.info(f"Fetching job description from URL: {url}")
            
            # Fetch content with timeout
            async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = await client.get(url, headers=headers)
                
                if response.status_code != 200:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"Failed to fetch URL: HTTP {response.status_code}"
                    )
                
                # Parse HTML content
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Extract job description text
                text = self._extract_job_text_from_html(soup)
                
                if not text or len(text.strip()) < 100:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail="Could not extract sufficient job description content from URL"
                    )
                
                # Try to extract company name
                company_name = self._extract_company_name(soup, url)
                
                logger.info(f"Successfully fetched job description from {url} ({len(text)} chars)")
                
                return {
                    'text': text,
                    'company_name': company_name
                }
                
        except httpx.TimeoutException:
            raise HTTPException(
                status_code=status.HTTP_408_REQUEST_TIMEOUT,
                detail="Request timeout while fetching URL"
            )
        except httpx.RequestError as e:
            logger.error(f"Error fetching URL {url}: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to fetch URL: {str(e)}"
            )
        except Exception as e:
            logger.error(f"Error processing URL {url}: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Error processing URL content"
            )
    
    def _is_valid_url(self, url: str) -> bool:
        """Validate URL format"""
        try:
            result = urlparse(url)
            return all([result.scheme in ['http', 'https'], result.netloc])
        except Exception:
            return False
    
    def _extract_job_text_from_html(self, soup: BeautifulSoup) -> str:
        """
        Extract job description text from HTML using various strategies
        """
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # Try to find job description in common containers
        job_selectors = [
            {'class': re.compile(r'job[-_]description', re.I)},
            {'class': re.compile(r'description', re.I)},
            {'class': re.compile(r'job[-_]details', re.I)},
            {'class': re.compile(r'content', re.I)},
            {'id': re.compile(r'job[-_]description', re.I)},
            {'id': re.compile(r'description', re.I)},
        ]
        
        for selector in job_selectors:
            element = soup.find(['div', 'section', 'article'], selector)
            if element:
                text = element.get_text(separator='\n', strip=True)
                if len(text) > 100:
                    return self._clean_text(text)
        
        # Fallback: get main content or body text
        main = soup.find(['main', 'article'])
        if main:
            text = main.get_text(separator='\n', strip=True)
            if len(text) > 100:
                return self._clean_text(text)
        
        # Last resort: get all body text
        body = soup.find('body')
        if body:
            text = body.get_text(separator='\n', strip=True)
            return self._clean_text(text)
        
        return soup.get_text(separator='\n', strip=True)
    
    def _extract_company_name(self, soup: BeautifulSoup, url: str) -> str:
        """
        Try to extract company name from HTML or URL
        """
        # Try meta tags
        company_meta = soup.find('meta', {'property': 'og:site_name'})
        if company_meta and company_meta.get('content'):
            return company_meta['content']
        
        # Try common class names
        company_selectors = [
            {'class': re.compile(r'company[-_]name', re.I)},
            {'class': re.compile(r'employer', re.I)},
        ]
        
        for selector in company_selectors:
            element = soup.find(['span', 'div', 'h1', 'h2'], selector)
            if element:
                company = element.get_text(strip=True)
                if company and len(company) < 100:
                    return company
        
        # Try to extract from URL domain
        try:
            domain = urlparse(url).netloc
            # Remove www. and common TLDs
            company = domain.replace('www.', '').split('.')[0]
            return company.capitalize()
        except Exception:
            pass
        
        return "Company from URL"
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        text = text.strip()
        return text
